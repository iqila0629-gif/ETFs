"""Build boss-facing v4 robustness report as a professional Word document (v6)."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY_DARK = RGBColor(0x33, 0x33, 0x33)
HEI = "SimHei"
YAHEI = "Microsoft YaHei"


def set_run_font(run, font_name, ea_font_name, size=None, bold=False, color=BLACK):
    run.font.name = font_name
    run.font.size = Pt(size) if size else run.font.size
    run.font.bold = bold
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), ea_font_name)


def make_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run_font(run, HEI, HEI, size={1: 16, 2: 14, 3: 12}.get(level, 11), bold=True, color=BLACK)
    return p


def add_normal(doc, text, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, YAHEI, YAHEI, size=size, color=BLACK)
    return p


def shade_cell(cell, hex_color):
    tc = cell._element.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tc.append(shading)


def add_table(doc, headers, rows, col_widths, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = w
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(header)
        set_run_font(r, HEI, HEI, size=font_size + 1, bold=True, color=BLACK)
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            bg = "E6E6E6" if row_idx % 2 == 1 else "FFFFFF"
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(value))
            set_run_font(r, YAHEI, YAHEI, size=font_size, color=GRAY_DARK)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
        '</w:tblBorders>'
    )
    table._tbl.tblPr.append(borders)
    return table


def dist_table_rows(data):
    rows = []
    for metric, full, test in data:
        rows.append([metric, "全历史"] + full)
        rows.append([metric, "2025后"] + test)
    return rows


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    normal = doc.styles["Normal"]
    normal.font.name = YAHEI
    normal.font.size = Pt(10)
    normal.font.color.rgb = BLACK
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), YAHEI)

    make_heading(doc, "ProFunds 策略稳健性与改善评估报告", 1)
    add_normal(doc, "日期：2026-08-18")
    add_normal(doc, "口径：m30_v2 正式版，103 支非货币基金、231 条策略；mapping 与 Excel、全部分析使用同一份当前 8/4 数据，mapping 与重算差异为 0。条件表述与正式 Excel 说明一致。")

    # 一、样本外有效性概览
    make_heading(doc, "一、样本外有效性概览", 2)
    add_normal(doc, "固定切分将历史数据分为建模段（train）与回测段（test）。建模段为 2005-06-15 ~ 2024-12-31；回测段为 2025-01-01 ~ 2026-08-04。严格重选指只用建模段数据重新选择策略，再用未参与建模的回测段数据回测，避免选策略时看过回测成绩。")

    make_heading(doc, "1.1 现役 m30_v2 固定切分", 3)
    add_table(doc, ["指标", "口径", "min", "P10", "中位数", "P90", "max"], [
        ["命中率", "建模段", "48.00%", "54.01%", "57.28%", "60.64%", "67.13%"],
        ["命中率", "回测段", "55.06%", "57.20%", "61.59%", "66.53%", "75.51%"],
        ["平均回报", "建模段", "0.139%", "0.221%", "0.321%", "0.424%", "0.566%"],
        ["平均回报", "回测段", "0.213%", "0.323%", "0.478%", "0.773%", "1.177%"],
        ["std", "建模段", "1.21%", "1.59%", "2.27%", "3.31%", "5.21%"],
        ["std", "回测段", "0.84%", "1.47%", "2.17%", "3.08%", "3.83%"],
        ["交易数", "建模段", "51", "121", "577", "1391", "2251"],
        ["交易数", "回测段", "33", "38", "95", "175", "235"],
    ], [Cm(2.4), Cm(1.9), Cm(1.8), Cm(1.8), Cm(2.0), Cm(1.8), Cm(1.8)])
    add_normal(doc, "结论：103/103 支基金回测段命中率 >=55%、平均回报 >=0.2%，没有基金在回测段掉出合格线，平均回报与命中率均高于建模段，未见样本外衰减。")

    make_heading(doc, "1.2 严格重选", 3)
    add_table(doc, ["指标", "口径", "min", "P10", "中位数", "P90", "max"], [
        ["命中率", "建模段", "48.00%", "54.71%", "57.20%", "61.85%", "67.13%"],
        ["命中率", "回测段", "44.00%", "55.16%", "60.47%", "65.97%", "75.61%"],
        ["平均回报", "建模段", "0.185%", "0.220%", "0.336%", "0.431%", "2.128%"],
        ["平均回报", "回测段", "-0.207%", "0.240%", "0.402%", "0.653%", "1.031%"],
        ["std", "建模段", "1.17%", "1.46%", "2.26%", "3.61%", "10.00%"],
        ["std", "回测段", "0.75%", "1.24%", "2.04%", "3.20%", "4.68%"],
        ["交易数", "建模段", "27", "103", "577", "1352", "2419"],
        ["交易数", "回测段", "30", "36", "76", "151", "206"],
    ], [Cm(2.4), Cm(1.9), Cm(1.8), Cm(1.8), Cm(2.0), Cm(1.8), Cm(1.8)])
    add_normal(doc, "结论：严格重选后仍有 99/103 支命中率 >=50%、93/103 支 >=55%，策略效果不是只靠看过回测段才成立。")

    make_heading(doc, "1.3 训练效果异常与改进", 3)
    add_normal(doc, "需要从候选池补强或替换（同时展示全历史与 2025 后）：")
    add_table(doc, ["基金", "条件", "全历史平均", "全历史命中", "2025后平均", "2025后命中", "问题"], [

        ["UKPIX", "XLC<0 且 信用利差>0，买基金后2日平均回报", "0.303%", "55.2%", "0.286%", "61.5%", "训练效果偏弱"],
        ["UTPSX", "GLD>0 且 JNK<0 且 XLU<0，买基金第二天", "0.236%", "55.9%", "0.396%", "59.6%", "训练效果偏弱"],
    ], [Cm(1.8), Cm(4.6), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8), Cm(2.2)], font_size=8)
    add_normal(doc, "全历史偏弱但 2025 后明显改善，暂不删除、继续观察：")
    add_table(doc, ["基金", "条件", "全历史平均", "全历史命中", "2025后平均", "2025后命中", "备注"], [
        ["UDPSX", "QQQ<=-1% 且 VIX5d>0，买基金第二天", "0.205%", "57.3%", "0.536%", "67.9%", "全历史弱且波动高，2025后改善明显，按观察处理"],
        ["UIPSX", "FXY<0 且 GLD>0 且 XLC>0，买基金后2日平均回报", "0.235%", "55.6%", "0.526%", "60.5%", "全历史弱，2025后改善明显，按观察处理"],
        ["CYPSX", "EEM<0 且 SLV>0 且 XLF<0，买基金第二天", "0.218%", "59.1%", "0.483%", "65.9%", "2025后明显改善"],
        ["IDPSX", "SPY<0 且 美元/黄金<=-1.0%，买基金第二天", "0.248%", "58.8%", "0.606%", "62.8%", "2025后明显改善"],
        ["OEPIX", "XLU<=-1% 且 VIX收盘>0，买基金后3日平均回报", "0.317%", "55.0%", "0.610%", "65.4%", "2025后明显改善"],
        ["SGPIX", "GDX>0 且 XLC<0 且 XLK<0，买基金第二天", "0.211%", "55.2%", "0.605%", "66.7%", "2025后明显改善"],
        ["SVPIX", "SLV>0 且 SPY<0 且 TIP>0，买基金第二天", "0.205%", "55.7%", "0.596%", "67.5%", "2025后明显改善"],
        ["UKPIX", "LQD<0 且 TIP<0 且 XLC<0，买基金后2日平均回报", "0.247%", "57.5%", "0.405%", "64.3%", "2025后明显改善"],
    ], [Cm(1.8), Cm(4.6), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8)], font_size=8)
    add_normal(doc, "说明：UDPSX、UIPSX 已移入观察名单。它们 2025 后平均回报 0.536%/0.526%、命中率 67.9%/60.5%，明显改善，但全历史平均回报与显著性仍偏弱，UDPSX 同时波动偏高，因此不直接删除，继续观察。")
    add_normal(doc, "改进措施：优先用候选池中更显著的信号替换 UKPIX、UTPSX 两条训练弱策略；全历史弱但 2025 后明显的 8 条继续观察，不为了凑覆盖硬塞低质量信号。")

    # 二、策略解释概览
    make_heading(doc, "二、策略解释概览", 2)
    add_normal(doc, "231 条策略已逐条给出经济解释，输出列包括“触发条件”“条件解释”与“机制”，不含条件代码和历史统计；全部明细见 v4_strategy_explanation_v3.csv。机制列按“条件说明了什么 → 带来的市场影响 → 基金产品特征 → 为什么该基金上涨/下跌”四段式书写。")

    make_heading(doc, "2.1 反向基金方向冲突（需改善）", 3)
    add_normal(doc, "反向基金应在标的下跌时上涨，但以下策略在“标的上涨/风险偏好回暖”条件下买入反向基金，等于赌风险偏好次日立即反转，经济逻辑弱：")
    add_table(doc, ["基金", "条件", "平均回报", "命中率", "问题"], [
        ["BITIX", "EEM>0 且 JNK>0 且 SPY>0，买基金第二天", "0.312%", "55.9%", "反向基金使用风险偏好条件"],
        ["BITIX", "EEM>0 且 GLD>0 且 SLV>0，买基金后2日平均回报", "0.346%", "55.0%", "反向基金使用风险偏好+贵金属条件"],
        ["SNPIX", "QQQ>0 且 XLC>0 且 XLV>0，买基金第二天", "0.371%", "57.0%", "反向基金使用风险偏好条件"],
        ["SNPSX", "SPY>0 且 XLC>0 且 XLK>0，买基金第二天", "0.421%", "61.5%", "反向基金使用风险偏好条件"],
        ["UCPIX", "JNK>0 且 LQD<0 且 XLC>0，买基金第二天", "0.319%", "55.3%", "反向基金信用信号混乱"],
        ["UHPIX", "HYG>0 且 SLV>0 且 XLC>0，买基金第二天", "0.428%", "57.3%", "反向基金使用风险偏好+贵金属条件"],
        ["UHPSX", "FXY<0 且 IWM>0 且 XLE<0，买基金第二天", "0.564%", "57.9%", "反向基金条件方向不明"],
        ["UKPSX", "QQQ>0 且 SPY>0 且 TIP<0，买基金第二天", "0.390%", "58.6%", "反向基金使用风险偏好条件"],
        ["USPIX", "EEM>0 且 SLV>0 且 XLV>0，买基金第二天", "0.497%", "56.6%", "反向基金使用风险偏好条件"],
        ["USPSX", "EEM>0 且 GLD>0 且 HYG>0，买基金第二天", "0.487%", "55.4%", "反向基金使用风险偏好+避险条件"],
        ["UVPIX", "EEM>0 且 UUP>0 且 XLF<0，买基金第二天", "0.300%", "58.3%", "反向基金使用新兴市场上涨条件"],
    ], [Cm(1.7), Cm(4.6), Cm(1.7), Cm(1.7), Cm(3.3)], font_size=8)

    make_heading(doc, "2.2 比特币基金解释受限（情有可原）", 3)
    add_table(doc, ["基金", "条件", "平均回报", "命中率", "原因"], [
        ["BITIX", "XLC>0 且 该基金当日回报<0，买基金第二天", "0.565%", "57.9%", "无比特币参考ETF，只能依赖代理变量"],
        ["BTCFX", "QQQ<0 且 XLK<0 且 XLV<0，买基金第二天", "0.687%", "57.1%", "无比特币参考ETF，只能依赖代理变量"],
        ["BTCFX", "EEM<0 且 信用利差<0，买基金后2日平均回报", "0.414%", "56.4%", "无比特币参考ETF，只能依赖代理变量"],
        ["BTCFX", "EEM<0 且 IWM<0 且 SPY<0，买基金后2日平均回报", "0.461%", "56.6%", "无比特币参考ETF，只能依赖代理变量"],
    ], [Cm(1.7), Cm(4.6), Cm(1.7), Cm(1.7), Cm(3.3)], font_size=8)
    add_normal(doc, "改进措施：反向基金优先替换为方向一致的条件，例如做空新兴市场用 EEM<0、做空科技用 QQQ<0 或 XLK<0；若保留，须明确标注为均值回归逻辑并做样本外验证。比特币基金暂不处理，未来加入加密参考资产后优先重做。")

    make_heading(doc, "2.3 同基金/同主题解释一致性检查", 3)
    add_normal(doc, "按“基金”维度检查：没有同一基金在相同市场状态（风险偏好收缩/回暖）下同时出现“价格倾向上涨”和“价格倾向承压”的矛盾解释。按“主题”维度存在方向相反的结论，但原因是主题内同时包含多头与反向基金，例如中国主题既有 Ultra China 多头、也有 Ultra Short China 反向，风险偏好收缩时多头承压、反向产品上涨，属于产品属性，不是解释矛盾。")
    add_normal(doc, "示例（银行金融类）：BKPIX/BKPSX 的两条解释分别是“风险偏好收缩时价格倾向承压”和“风险偏好回暖时价格倾向上涨”，属于不同市场状态下的互补结论，不构成矛盾。检查明细见 v4_explanation_conflict_check.csv。")
    add_normal(doc, "解释缺口检查：原机制对“承压却买入”的解释存在逻辑缺口（136/231 条）。已改为更合理的“风险冲击/防御落后后的次日均值回归”解释；修订后仅剩 7/231 条仍写“承压”，且全部是反向基金使用风险偏好条件的已知方向冲突问题。明细见 v4_explanation_weak_pressure.csv。")

    # 三、风险概览
    make_heading(doc, "三、风险概览", 2)

    make_heading(doc, "3.1 合并策略风险分布与波动比值（103 基金）", 3)
    add_normal(doc, "合并策略相对基金无条件波动比值：min 0.491，P10 0.686，中位数 1.023，P90 1.277，max 1.535。")
    add_table(doc, ["比值区间", "基金数", "占比"], [
        ["<0.5x", "1", "0.97%"],
        ["0.5-1.0x", "46", "44.66%"],
        ["1.0-1.5x", "54", "52.43%"],
        ["1.5-2.0x", "2", "1.94%"],
        [">2.0x", "0", "0.00%"],
    ], [Cm(4.0), Cm(3.5), Cm(3.5)])
    add_normal(doc, "详细指标分布：")
    merged_rows = dist_table_rows([
        ("平均回报", ["0.207%", "0.247%", "0.361%", "0.445%", "0.523%"], ["0.213%", "0.323%", "0.478%", "0.773%", "1.177%"]),
        ("命中率", ["53.79%", "55.38%", "57.93%", "61.07%", "66.32%"], ["55.06%", "57.20%", "61.59%", "66.53%", "75.51%"]),
        ("std", ["1.25%", "1.57%", "2.27%", "3.22%", "4.46%"], ["0.84%", "1.47%", "2.17%", "3.08%", "3.83%"]),
        ("年化波动", ["19.78%", "24.90%", "36.09%", "51.11%", "70.85%"], ["13.35%", "23.31%", "34.42%", "48.91%", "60.72%"]),
        ("下行标准差", ["0.71%", "0.85%", "1.36%", "2.00%", "2.98%"], ["0.40%", "0.55%", "1.03%", "1.68%", "2.21%"]),
        ("平均/std", ["0.076", "0.107", "0.156", "0.224", "0.270"], ["0.109", "0.163", "0.248", "0.355", "0.406"]),
        ("t 统计量", ["1.05", "2.41", "3.92", "5.63", "6.51"], ["0.65", "1.57", "2.31", "3.11", "4.18"]),
    ])
    add_table(doc, ["指标", "口径", "min", "P10", "中位数", "P90", "max"], merged_rows,
              [Cm(2.4), Cm(1.9), Cm(1.8), Cm(1.8), Cm(2.0), Cm(1.8), Cm(1.8)], font_size=8)

    make_heading(doc, "3.2 单策略风险分布（231 条）", 3)
    add_normal(doc, "单策略相对基金无条件波动比值：min 0.442，P10 0.635，中位数 1.026，P90 1.328，max 1.655。")
    add_table(doc, ["比值区间", "策略数", "占比"], [
        ["<0.5x", "4", "1.73%"],
        ["0.5-1.0x", "107", "46.32%"],
        ["1.0-1.5x", "111", "48.05%"],
        ["1.5-2.0x", "9", "3.90%"],
        [">2.0x", "0", "0.00%"],
    ], [Cm(4.0), Cm(3.5), Cm(3.5)])
    add_normal(doc, "详细指标分布：")
    single_rows = dist_table_rows([
        ("平均回报", ["0.202%", "0.255%", "0.358%", "0.520%", "0.952%"], ["0.204%", "0.286%", "0.555%", "0.938%", "1.827%"]),
        ("命中率", ["54.97%", "55.43%", "58.00%", "62.20%", "66.32%"], ["53.97%", "56.41%", "62.22%", "68.33%", "80.49%"]),
        ("std", ["0.82%", "1.56%", "2.21%", "3.73%", "5.04%"], ["0.79%", "1.32%", "2.19%", "3.53%", "5.11%"]),
        ("年化波动", ["13.03%", "24.79%", "35.13%", "59.14%", "79.93%"], ["12.56%", "20.88%", "34.84%", "56.12%", "81.13%"]),
        ("下行标准差", ["0.41%", "0.85%", "1.36%", "2.30%", "3.50%"], ["0.21%", "0.60%", "1.06%", "2.06%", "3.08%"]),
        ("平均/std", ["0.057", "0.104", "0.161", "0.225", "0.359"], ["0.091", "0.142", "0.270", "0.391", "0.748"]),
        ("t 统计量", ["1.05", "2.13", "3.22", "4.17", "5.16"], ["0.59", "1.07", "1.93", "2.84", "4.93"]),
    ])
    add_table(doc, ["指标", "口径", "min", "P10", "中位数", "P90", "max"], single_rows,
              [Cm(2.4), Cm(1.9), Cm(1.8), Cm(1.8), Cm(2.0), Cm(1.8), Cm(1.8)], font_size=8)

    make_heading(doc, "3.3 风险异常与改进", 3)
    add_normal(doc, "产品命名说明：UltraSector = 2 倍杠杆行业基金；Ultra China/Japan/Latin America 等 = 2 倍杠杆地区/主题基金；UltraShort/Short = 反向（做空）产品；Bull/Bear = 多头/空头；带 Ultra 前缀一般表示约 2 倍杠杆。")
    add_normal(doc, "需观察基金（合并层面，含比值 >=1.5）：")
    add_table(doc, ["基金", "关注点", "数值", "说明"], [
        ["BLPIX", "合并 std 为基金无条件 std 的 1.54 倍", "1.80% vs 1.17%", "杠杆多头，波动高于基金本身"],
        ["BLPSX", "合并 std 为基金无条件 std 的 1.52 倍", "1.79% vs 1.18%", "杠杆多头，波动高于基金本身"],
        ["SPPIX", "2025 后 std 高于建模段 1.5 倍", "3.05% vs 2.02%", "反向贵金属，波动上升需观察"],
        ["SPPSX", "2025 后 std 高于建模段 1.5 倍", "3.05% vs 2.02%", "反向贵金属，波动上升需观察"],
        ["UXPSX", "2025 后 std 高于建模段 1.5 倍", "1.95% vs 1.21%", "反向国际，波动上升需观察"],
    ], [Cm(1.8), Cm(4.5), Cm(2.8), Cm(3.8)], font_size=8)
    add_normal(doc, "情有可原（先写产品，再写指标；只列异常的指标）：")
    add_table(doc, ["基金", "关注点", "备注"], [
        ["PMPIX / PMPSX", "策略 std 4.64% / 4.63%", "贵金属 UltraSector 杠杆产品，基金本身全历史 std 3.53% / 3.59%"],
        ["UHPIX", "策略 std 4.46%", "Ultra Short China 反向杠杆产品，基金本身全历史 std 4.17%"],
        ["UBPIX / UBPSX", "最大单笔亏损 -40.5%", "Ultra Latin America 杠杆产品，基金本身最大单笔亏损 -40.5%"],
        ["UAPIX / UAPSX", "最大单笔亏损 -32.4%", "Ultra Small-Cap 杠杆产品，基金本身最大单笔亏损 -50.0% / -41.5%"],
        ["UGPSX", "2025 后 std 5.11%", "Ultra China 杠杆产品，基金本身全历史 std 4.21%"],
        ["SMPIX / SMPSX / UGPIX", "2025 后 std 4.48%-4.90%", "UltraSector/Ultra 杠杆产品，基金本身全历史 std 3.13% / 3.14% / 4.23%"],
    ], [Cm(3.0), Cm(4.0), Cm(7.0)], font_size=8)
    add_normal(doc, "改进措施：SPPIX/SPPSX/UXPSX 列入观察名单，若波动继续上升再评估增加风险门槛；BLPIX/BLPSX 因合并比值 >=1.5 保留在需观察；其余杠杆/反向产品波动高是基金本身特质，不删除。")

    # 四、阈值单调性概览
    make_heading(doc, "四、阈值单调性概览", 2)
    add_normal(doc, "本章假设：对单个策略，把某个条件阈值改深，例如 XLK<=-1% 且 该基金当日回报<=-2% 中的 XLK 改为 <=-2%/-2.5%/-3%，整个策略的平均回报和命中率应递增。以下全部按策略统计；多条件策略只调整被扫描的那个条件，其他子条件保持不变。4.1、4.2、4.3 使用同一份策略级阈值扫描结果。扫描范围已扩展到全部数值条件：包括 ETF/自身的 up/down（0 阈值）、big/gt/lt、外部 up/down/ge/le，以及连跌天数；以交付阈值为中心，同时扫描更松和更紧的阈值（例如原 VIX>2，则 1/1.5/2/2.5/3 都计算）。共 563 组扫描（562 组数值阈值 + 1 组连跌天数），全部进入方向一致性统计。")

    make_heading(doc, "4.1 策略阈值扫描汇总（以交付阈值为中心，偏移为正表示更严格）", 3)
    add_normal(doc, "说明：偏移加在“阈值的绝对值”上，正偏移表示更严格，对上涨和下跌条件都一样。例如 EEM>0 的偏移 +1 表示 EEM>1%；QQQ<0 的偏移 +1 表示 QQQ<-1%；偏移 0 即交付阈值本身。对 EEM>0、JNK<0 这类交付阈值为 0 的条件，只有正偏移（更严格）存在；负偏移会让幅度变成负数，已被过滤。因此 +0.5% 行策略数多，是因为几乎所有 0 阈值条件都会贡献；负偏移行策略数少，是因为只有交付阈值本身大于 0.5 的策略（如 VIX5d>=10%）才有。某偏移下策略数偏少，还因为该偏移对应阈值下部分策略无法触发（有效交易数为 0 或数据不足），只有能触发且有数据的策略计入。")
    add_normal(doc, "负向阈值扫描（策略级，其余条件固定；同一策略调整不同条件计为不同策略）：")
    add_table(doc, ["阈值偏移", "平均回报中位数", "命中率中位数", "交易数中位数", "策略数"], [
        ["-2%", "0.147%", "52.5%", "924", "8"],
        ["-1.5%", "0.249%", "55.3%", "1193", "9"],
        ["-1%", "0.177%", "55.5%", "1051", "32"],
        ["-0.5%", "0.232%", "55.9%", "817", "37"],
        ["0%（交付阈值）", "0.350%", "57.8%", "411", "331"],
        ["+0.5%", "0.365%", "58.1%", "243", "323"],
        ["+1%", "0.443%", "58.3%", "118", "286"],
        ["+1.5%", "0.470%", "59.3%", "67", "214"],
        ["+2%", "0.490%", "59.6%", "64", "118"],
    ], [Cm(2.6), Cm(3.0), Cm(3.0), Cm(3.0), Cm(2.0)], font_size=8)
    add_normal(doc, "正向阈值扫描（策略级，其余条件固定）：")
    add_table(doc, ["阈值偏移", "平均回报中位数", "命中率中位数", "交易数中位数", "策略数"], [
        ["-2%", "0.300%", "56.4%", "693", "11"],
        ["-1.5%", "0.297%", "56.7%", "653", "11"],
        ["-1%", "0.266%", "56.6%", "612", "18"],
        ["-0.5%", "0.260%", "57.5%", "571", "18"],
        ["0%（交付阈值）", "0.343%", "57.9%", "374", "265"],
        ["+0.5%", "0.324%", "58.1%", "208", "238"],
        ["+1%", "0.334%", "58.9%", "135", "189"],
        ["+1.5%", "0.355%", "59.4%", "149", "141"],
        ["+2%", "0.333%", "57.3%", "106", "98"],
    ], [Cm(2.6), Cm(3.0), Cm(3.0), Cm(3.0), Cm(2.0)], font_size=8)
    add_normal(doc, "结论：负向策略在更严格一侧平均回报和命中率中位数整体仍小幅上升，但交易数明显下降，交付阈值是平衡点；正向策略在交付阈值附近平均回报最高，更严格后命中率略升、交易数下降。")

    make_heading(doc, "4.2 策略连跌扫描（策略级）", 3)
    add_normal(doc, "现役策略中只有 GVPSX 使用 self_3down（长期国债下跌 且 该基金连续3日下跌），按策略扫描连续下跌天数，其他条件固定：")
    add_table(doc, ["连续下跌天数", "平均回报", "命中率", "交易数"], [
        ["1 天", "-0.029%", "50.2%", "793"],
        ["2 天", "+0.097%", "54.3%", "1023"],
        ["3 天", "+0.207%", "58.0%", "400"],
        ["4 天", "+0.278%", "59.6%", "89"],
        ["5 天", "无触发", "无", "0"],
    ], [Cm(3.0), Cm(3.6), Cm(3.0), Cm(3.0)], font_size=8)
    add_normal(doc, "结论：连跌天数从 1 增至 4，平均回报与命中率单调上升，符合“连跌越久反弹越强”的假设；第 5 天叠加外部条件后无触发。")

    make_heading(doc, "4.3 方向相反的异常策略", 3)
    add_normal(doc, "按策略逐条统计阈值加深后的效果方向，同时看平均回报与命中率两个相关系数；任一指标相关系数 < -0.5 即列入异常。全部 563 组策略级阈值扫描中，负向策略平均回报中位数相关系数为 +0.76、命中率中位数为 +0.63；正向策略分别为 +0.45、+0.02。初筛 209 组异常；其中“触发阈值太少”且平均回报呈正向单调的 22 组已剔除（样本少但方向与预期一致，不应判为异常），最终保留 187 组。")
    import pandas as _pd
    _cat = _pd.read_csv(r"C:\Users\vanessacen\Desktop\新基金预测\04_结果\v4_中间结果\v4_稳健性分析\v4_strategy_direction_category_summary.csv")
    _desc = {
        "末端反转": "平均回报先升后降，最后 1-2 个阈值明显回落",
        "触发阈值太少": "有效阈值点不足 4 个，无法可靠判断趋势",
        "忽上忽下": "阈值间方向频繁切换，无稳定趋势",
        "整体反向单调": "阈值越严格，平均回报持续下降",
        "拐点与整体不同": "存在与整体方向不一致的拐点",
    }
    add_normal(doc, "异常类别概览：")
    add_table(doc, ["异常类别", "数量", "占比", "特征"], [
        [r.category, str(r.count), f"{r.pct:.1f}%", _desc[r.category]]
        for r in _cat.itertuples(index=False)
    ], [Cm(3.0), Cm(2.0), Cm(2.0), Cm(7.0)], font_size=8)
    _rows = _pd.read_csv(r"C:\Users\vanessacen\Desktop\新基金预测\04_结果\v4_中间结果\v4_稳健性分析\v4_strategy_direction_category_rows.csv").sort_values(["category", "rho_avg"])
    add_normal(doc, "异常策略明细（按异常类型分组排序）：")
    _detail = [[r.ticker, r.condition_text, r.scan_token, f"{r.rho_avg:.2f}", f"{r.rho_hit:.2f}", r.flag, r.category] for r in _rows.itertuples(index=False)]
    add_table(doc, ["基金", "条件", "调整条件", "平均回报rho", "命中率rho", "异常指标", "异常类型"], _detail,
              [Cm(1.5), Cm(4.0), Cm(2.0), Cm(1.6), Cm(1.6), Cm(2.0), Cm(2.2)], font_size=7)
    add_normal(doc, "逐策略阈值明细改放 CSV：v4_strategy_direction_detail_v2.csv。改进措施：对这些异常策略不做“更深阈值”外推，优先保留交付阈值。")

    out = r"C:\Users\vanessacen\Desktop\新基金预测\03_文档\报告\2026-08-18_v4稳健性分析报告_上司版_修改版.docx"
    doc.save(out)
    print("OK: " + out)


if __name__ == "__main__":
    main()





